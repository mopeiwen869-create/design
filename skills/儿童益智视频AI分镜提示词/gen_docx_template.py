"""
丁桂儿分镜脚本 Word 文档生成器（紧凑文字版）
所有输出使用紧凑文字排版，禁止表格。
skill 调用此脚本生成标准格式的 .docx 文件

用法: python gen_docx_template.py <data_json> <output_path>
  data_json: JSON 文件路径，包含所有分镜数据
  output_path: 输出 .docx 路径

data_json 格式:
{
  "topic": "选题名称",
  "duration": "40秒",
  "tool": "可灵3.0",
  "style": "Q版2D卡通亲子科普",
  "atmosphere": { ... },
  "settings": { ... },
  "shots": [ { "title": "分镜1", "time": "0-6s", "景别": "...", "构图": "...", "运镜": "...", "转场": "...", "画面": "...", "旁白": "...", "花字": "...", "音效": "...", "埋点": "...", "后期": "..." }, ... ],
  "voiceover": [ { "time": "0-6秒", "text": "..." }, ... ],
  "verification": [ { "item": "...", "source": "..." }, ... ],
  "timeline": [ { "time": "...", "visual": "...", "audio": "...", "caption": "...", "sfx": "..." }, ... ],
  "caption_list": [ { "in": "...", "out": "...", "content": "...", "anim": "...", "style": "...", "note": "..." }, ... ],
  "assets": [ { "id": "A01", "shot": "...", "duration": "...", "elements": "...", "prompt": "...", "status": "⬜" }, ... ],
  "checklist": [ "...", ... ]
}
"""

import json, sys, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def build_doc(data):
    doc = Document()

    # 默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    rpr = style.element.rPr
    if rpr is None:
        from docx.oxml import OxmlElement
        rpr = OxmlElement('w:rPr')
        style.element.append(rpr)
    rFonts = rpr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts')
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement('w:rFonts')
        rpr.append(rFonts)
    rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')

    dt = data

    # ── 辅助函数 ──
    def h1(text):
        h = doc.add_heading(text, level=1)
        for r in h.runs:
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
        return h

    def h2(text):
        h = doc.add_heading(text, level=2)
        for r in h.runs:
            r.font.name = '微软雅黑'
            r.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
        return h

    def p(text, bold=False, size=None, color=None, align=None):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
        if bold:
            run.bold = True
        if size:
            run.font.size = Pt(size)
        if color:
            run.font.color.rgb = RGBColor(*color)
        if align:
            para.alignment = align
        return para

    def rich_para(parts, spacing_after=None):
        """parts: [(text, bold, color_or_None), ...]"""
        para = doc.add_paragraph()
        for text, bold, color in parts:
            run = para.add_run(text)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '微软雅黑')
            run.bold = bold
            if color:
                run.font.color.rgb = RGBColor(*color)
        if spacing_after is not None:
            para.paragraph_format.space_after = Pt(spacing_after)
        return para

    def sep_line():
        """添加分隔线：---"""
        p('─' * 50, size=8, color=(180, 180, 180))

    # ═══ 封面 ═══
    p(dt.get('topic', ''), bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    p('')
    info_parts = []
    if 'duration' in dt:
        info_parts.append(f"时长：{dt['duration']}")
    if 'tool' in dt:
        info_parts.append(f"格式：{dt['tool']}")
    if 'style' in dt:
        info_parts.append(f"风格：{dt['style']}")
    p('  |  '.join(info_parts), size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    p('')
    doc.add_page_break()

    # ═══ 一、氛围与画质 ═══
    h1('一、氛围与画质')
    atm = dt.get('atmosphere', {})
    for key in ['风格核心', '视觉基调', '色彩与影调', '字幕呈现']:
        if key in atm:
            rich_para([(key + '：', True, None), (atm[key], False, None)])

    # ═══ 二、基础设定 ═══
    h1('二、基础设定')
    st = dt.get('settings', {})
    if 'character' in st:
        rich_para([('核心角色-丁桂儿：', True, None), (st['character'], False, None)])
    if 'emotions' in st and st['emotions']:
        p('丁桂儿表情系统（按场景调用）：', bold=True)
        for k, v in st['emotions'].items():
            p(f"  - {k} → {v}", size=10)
    for key in ['辅助角色', '场景', '声音']:
        if key in st:
            p(f"{key}：{st[key]}")
    doc.add_page_break()

    # ═══ 三、分镜脚本（紧凑文字版） ═══
    h1('三、分镜脚本')
    for shot in dt.get('shots', []):
        # 构建紧凑标题行
        title = shot.get('title', '')
        time = shot.get('time', '')
        scene = shot.get('景别', '')
        comp = shot.get('构图', '')
        cam = shot.get('运镜', '')
        trans = shot.get('转场', '')
        header = f"【{title}｜{time}｜{scene}｜{comp}｜{cam}｜{trans}】"
        p(header, bold=True, size=10)

        # 画面
        if '画面' in shot:
            p(f"画面：{shot['画面']}", size=9.5)
        # 旁白/花字/音效/埋点/后期
        for field, label in [('旁白', '旁白｜'), ('花字', '花字｜'), ('音效', '【音效】'), ('埋点', '【埋点】'), ('后期', '【后期】')]:
            if field in shot and shot[field]:
                p(f"{label}{shot[field]}", size=9.5)
        p('')  # 分镜之间空行

    doc.add_page_break()

    # ═══ 四、口播文字 ═══
    h1('四、口播文字（逐字稿）')
    vo = dt.get('voiceover', [])
    for item in vo:
        rich_para([
            ('▶ ' + item.get('time', '') + '  ', True, (192, 0, 0)),
            (item.get('text', ''), False, None)
        ])
        p('')

    # ═══ 五、正确性校验表 ═══
    h1('五、正确性校验表')
    p('每条关键信息的医学依据来源核验：')
    vf = dt.get('verification', [])
    for v in vf:
        sep_line()
        p(f"口播内容：{v.get('item', '')}", size=10)
        p(f"依据：{v.get('source', '')}", size=9, color=(80, 80, 80))
    if vf:
        sep_line()

    # ═══ 六、剪辑对照表 ═══
    doc.add_page_break()
    h1('六、剪辑对照表 — 四轨时间轴总览')
    tl = dt.get('timeline', [])
    for t in tl:
        sep_line()
        time_str = t.get('time', '')
        p(f"【{time_str}】", bold=True, size=10)
        p(f"画面：{t.get('visual', '')}", size=9.5)
        audio = t.get('audio', '')
        if audio:
            p(f"旁白：{audio}", size=9.5)
        caption = t.get('caption', '')
        if caption:
            p(f"花字：{caption}", size=9.5)
        sfx = t.get('sfx', '')
        if sfx:
            p(f"音效：{sfx}", size=9.5)
    if tl:
        sep_line()

    # ═══ 七、花字包装明细表 ═══
    h1('七、花字包装明细表')
    cl = dt.get('caption_list', [])
    for c in cl:
        sep_line()
        p(f"【{c.get('in', '')}-{c.get('out', '')}】", bold=True, size=10)
        p(f"内容：{c.get('content', '')}", size=9.5)
        p(f"动画：{c.get('anim', '')}", size=9.5)
        p(f"颜色/样式：{c.get('style', '')}", size=9.5)
        note = c.get('note', '')
        if note:
            p(f"备注：{note}", size=9, color=(80, 80, 80))
    if cl:
        sep_line()

    # ═══ 八、AI素材生成清单 ═══
    h1('八、AI素材生成清单')
    as_ = dt.get('assets', [])
    for a in as_:
        sep_line()
        p(f"{a.get('id', '')} | {a.get('shot', '')} | {a.get('duration', '')}", bold=True, size=10)
        p(f"关键元素：{a.get('elements', '')}", size=9.5)
        p(f"生成提示词摘要：{a.get('prompt', '')}", size=9.5)
        p(f"状态：{a.get('status', '⬜')}", size=9.5)
    if as_:
        sep_line()
        p('')
    p('💡 建议：所有AI素材按编号批量生成后再入时间线剪辑，不要逐段生成逐段剪。', bold=True, size=10)

    # ═══ 九、输出自检清单 ═══
    doc.add_page_break()
    h1('九、输出自检清单')
    for item in dt.get('checklist', []):
        p('✅ ' + item, size=10)

    return doc


def main():
    if len(sys.argv) < 3:
        print("用法: python gen_docx_template.py <data_json> <output_path>")
        sys.exit(1)
    data_path = sys.argv[1]
    output_path = sys.argv[2]
    with open(data_path, 'r', encoding='utf-8') as f:
        data = json.load(f)
    doc = build_doc(data)
    doc.save(output_path)
    print(f"✅ Word已生成: {output_path}")


if __name__ == '__main__':
    main()
